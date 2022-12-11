# Reading and Writing MS Word Files in Python via Python-Docx Module
# https://stackabuse.com/reading-and-writing-ms-word-files-in-python-via-python-docx-module/

# Working With Text In Python .docx Module
# https://www.geeksforgeeks.org/working-with-text-in-python-docx-module/

import win32com.client

import os

import docx
from htmldocx import HtmlToDocx


def save_html_to_word2(data, filename):
    # https://www.blog.pythonlibrary.org/2010/07/16/python-and-microsoft-office-using-pywin32/
    # http://p-nand-q.com/python/win32com_and_office.html
    # https://fortes-arthur.medium.com/handling-with-doc-extension-with-python-b6491792311e

    # Create new Word Object
    word = win32com.client.Dispatch("Word.Application")

    # Word Application should`t be visible
    word.Visible = 0

    # Create new Document Object
    doc = word.Documents.Add()

    # Make some Setup to the Document:
    # doc.PageSetup.Orientation = 1
    # doc.PageSetup.LeftMargin = 20
    # doc.PageSetup.TopMargin = 20
    # doc.PageSetup.BottomMargin = 20
    # doc.PageSetup.RightMargin = 20
    # doc.Content.Font.Size = 11
    # doc.Content.Paragraphs.TabStops.Add(100)

    doc.Range(0, 0).InsertAfter(data)
    # doc.Content.MoveEnd
    # doc.Content.Text = 'sdfsadfasdfsafdasdfsfd'

    # doc.SaveAs(filename, FileFormat=0)

    # Close the Word Document (a save-Dialog pops up)
    doc.Close()

    # Close the Word Application
    word.Quit()


def save_html_to_word(html, filename):
    document = docx.Document()
    new_parser = HtmlToDocx()
    new_parser.paragraph_style = 'Normal'

    # Convert string
    # docx = new_parser.parse_html_string(html)

    new_parser.add_html_to_document(html, document)

    # do more stuff to document
    document.save(filename)


def append_html_to_word(html, filename):
    if os.path.isfile(filename):
        document = docx.Document(filename)
    else:
        document = docx.Document()

    new_parser = HtmlToDocx()
    new_parser.paragraph_style = 'Normal'

    # Convert string
    # docx = new_parser.parse_html_string(html)

    new_parser.add_html_to_document(html, document)

    # do more stuff to document
    document.save(filename)
